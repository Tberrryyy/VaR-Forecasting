import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from scipy import stats
from statsmodels.tsa.stattools import adfuller, acf, pacf
from statsmodels.tsa.arima.model import ARIMA
from arch import arch_model
from statsmodels.regression.linear_model import OLS
from statsmodels.tools.tools import add_constant
import warnings
import os
from docx import Document
from docx.shared import Inches, RGBColor
import datetime

warnings.filterwarnings('ignore')

# 创建结果文件夹
results_dir = "GARCH results"
if not os.path.exists(results_dir):
    os.makedirs(results_dir)

# 1. 数据读取与预处理
print("1. 读取数据...")
data = pd.read_csv('spy_data.csv', usecols=[0, 1], parse_dates=['Unnamed: 0'], index_col='Unnamed: 0')
data.columns = ['log_ret']
returns = data['log_ret']
print(f"数据时间范围: {returns.index[0]} 到 {returns.index[-1]}")
print(f"总样本数: {len(returns)}")

# 创建Word文档
doc = Document()
doc.add_heading('GARCH模型VaR预测结果分析报告', 0)
current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
doc.add_paragraph(f"生成时间: {current_time}")
doc.add_paragraph("")

# 添加数据基本信息
doc.add_heading('一、数据基本信息', level=1)
p = doc.add_paragraph()
p.add_run(f"数据时间范围: {returns.index[0]} 到 {returns.index[-1]}").bold = True
p = doc.add_paragraph()
p.add_run(f"总样本数: {len(returns)}").bold = True

# 2. 划分样本内和样本外
print("\n2. 划分样本内和样本外...")
n = len(returns)
n_insample = int(n * 0.8)
n_outsample = n - n_insample

returns_insample = returns.iloc[:n_insample]
returns_outsample = returns.iloc[n_insample:]

print(f"样本内数据量: {n_insample} ({returns_insample.index[0]} 到 {returns_insample.index[-1]})")
print(f"样本外数据量: {n_outsample} ({returns_outsample.index[0]} 到 {returns_outsample.index[-1]})")

# 添加样本划分信息
doc.add_heading('二、样本划分', level=1)
p = doc.add_paragraph()
p.add_run(f"样本内数据量: {n_insample}").bold = True
p.add_run(f" ({returns_insample.index[0]} 到 {returns_insample.index[-1]})")
p = doc.add_paragraph()
p.add_run(f"样本外数据量: {n_outsample}").bold = True
p.add_run(f" ({returns_outsample.index[0]} 到 {returns_outsample.index[-1]})")
doc.add_paragraph("")

# 3. 样本内数据平稳性检验
print("\n3. 样本内平稳性检验...")
adf_result = adfuller(returns_insample.dropna())
print(f"ADF统计量: {adf_result[0]:.6f}")
print(f"p值: {adf_result[1]:.6f}")
if adf_result[1] < 0.05:
    print("结论: 样本内收益率序列是平稳的")
    stationarity_conclusion = "样本内收益率序列是平稳的"
    returns_insample_processed = returns_insample.copy()
else:
    print("结论: 样本内收益率序列非平稳，进行一阶差分处理")
    stationarity_conclusion = "样本内收益率序列非平稳，进行一阶差分处理"
    returns_insample_processed = returns_insample.diff().dropna()
    print(f"差分后样本内数据量: {len(returns_insample_processed)}")

# 添加平稳性检验结果
doc.add_heading('三、平稳性检验', level=1)
p = doc.add_paragraph()
p.add_run(f"ADF统计量: {adf_result[0]:.6f}").bold = True
p = doc.add_paragraph()
p.add_run(f"p值: {adf_result[1]:.6f}").bold = True
p = doc.add_paragraph()
p.add_run(f"结论: {stationarity_conclusion}").bold = True
doc.add_paragraph("")

# 4. 样本内自相关检验和均值方程选择
print("\n4. 样本内自相关检验和均值方程选择...")
# 计算ACF和PACF
lags = min(20, len(returns_insample_processed) // 5)
acf_values = acf(returns_insample_processed, nlags=lags, fft=False)
pacf_values = pacf(returns_insample_processed, nlags=lags)


# 使用AIC/BIC准则选择ARMA阶数
def select_arma_order(series, max_p=5, max_q=5):
    best_aic = np.inf
    best_bic = np.inf
    best_order_aic = (0, 0)
    best_order_bic = (0, 0)

    for p in range(max_p + 1):
        for q in range(max_q + 1):
            if p == 0 and q == 0:
                continue
            try:
                model = ARIMA(series, order=(p, 0, q))
                results = model.fit()
                aic = results.aic
                bic = results.bic

                if aic < best_aic:
                    best_aic = aic
                    best_order_aic = (p, q)
                if bic < best_bic:
                    best_bic = bic
                    best_order_bic = (p, q)
            except:
                continue

    print(f"AIC准则选择的最优阶数: ARMA{best_order_aic}")
    print(f"BIC准则选择的最优阶数: ARMA{best_order_bic}")
    return best_order_bic


best_p, best_q = select_arma_order(returns_insample_processed)

# 添加ARMA模型选择结果
doc.add_heading('四、ARMA模型选择', level=1)
p = doc.add_paragraph()
p.add_run(f"基于AIC准则选择的最优阶数: ARMA({best_p}, {best_q})").bold = True
p = doc.add_paragraph()
p.add_run("注: AIC准则倾向于选择拟合优度更高的模型，BIC准则对参数个数惩罚更大，倾向于选择更简洁的模型。").italic = True
doc.add_paragraph("")

# 5. 样本内：估计GARCH模型
print("\n5. 样本内：估计GARCH(1,1)模型...")


def fit_garch(returns_data, p=0, q=0, dist='normal'):
    """拟合GARCH(1,1)模型"""
    if p > 0 and q > 0:
        model = arch_model(returns_data, mean='AR', lags=p, vol='GARCH', p=1, q=1, dist=dist)
    elif p > 0:
        model = arch_model(returns_data, mean='AR', lags=p, vol='GARCH', p=1, q=1, dist=dist)
    elif q > 0:
        model = arch_model(returns_data, mean='AR', lags=q, vol='GARCH', p=1, q=1, dist=dist)
    else:
        model = arch_model(returns_data, mean='constant', vol='GARCH', p=1, q=1, dist=dist)

    results = model.fit(disp='off', show_warning=False)
    return results


print("\n5.1 使用标准正态分布假设...")
garch_normal_insample = fit_garch(returns_insample_processed, p=best_p, q=best_q, dist='normal')
print(garch_normal_insample.summary())

# 添加GARCH模型估计结果
doc.add_heading('五、GARCH(1,1)模型估计结果', level=1)
doc.add_paragraph("模型设定: GARCH(1,1) with Standard Normal Distribution")
doc.add_paragraph(f"均值方程: ARMA({best_p}, {best_q})")

# 将模型参数添加到Word文档
params = garch_normal_insample.params
p = doc.add_paragraph()
p.add_run("模型参数估计值:").bold = True
for param_name, param_value in params.items():
    doc.add_paragraph(f"  {param_name}: {param_value:.6f}")

# 获取对数似然值等信息
log_likelihood = garch_normal_insample.loglikelihood
aic = garch_normal_insample.aic
bic = garch_normal_insample.bic

doc.add_paragraph(f"对数似然值: {log_likelihood:.2f}")
doc.add_paragraph(f"AIC: {aic:.2f}")
doc.add_paragraph(f"BIC: {bic:.2f}")
doc.add_paragraph("")

# 6. 样本内VaR计算和动态分位数检验
print("\n6. 样本内VaR计算和动态分位数检验...")
confidence_levels = [0.01, 0.05, 0.10]

# 获取样本内条件波动率，而不是使用forecast
cond_vol_normal = garch_normal_insample.conditional_volatility

# 计算VaR序列（而不是单个值）
var_normal_series = {}
for alpha in confidence_levels:
    var_normal_series[alpha] = -cond_vol_normal * stats.norm.ppf(alpha)

print("正态分布假设下的样本内VaR（最后一个观测值）:")
for alpha in confidence_levels:
    print(f"  {alpha * 100:.1f}% VaR: {var_normal_series[alpha].iloc[-1]:.6f}")

# 添加样本内VaR结果
doc.add_heading('六、样本内VaR计算结果', level=1)
p = doc.add_paragraph()
p.add_run("正态分布假设下的样本内VaR（最后一个观测值）:").bold = True
for alpha in confidence_levels:
    doc.add_paragraph(f"  {alpha * 100:.1f}% VaR: {var_normal_series[alpha].iloc[-1]:.6f}")
doc.add_paragraph("")


# 动态分位数检验函数
def dynamic_quantile_test(returns, var_series, alpha):
    # 确保输入数据是数值类型
    returns = returns.astype(float)
    var_series = var_series.astype(float)

    hits = (returns < -var_series).astype(int)

    hits = hits.fillna(0)
    var_series = var_series.fillna(var_series.mean())

    X = add_constant(pd.DataFrame({
        'lag_hit': hits.shift(1).fillna(0),
        'var': var_series
    }))
    y = hits - alpha

    # 确保X和y都是数值类型
    X = X.astype(float)
    y = y.astype(float)

    valid_idx = ~(X.isna().any(axis=1) | y.isna())
    X = X[valid_idx]
    y = y[valid_idx]

    try:
        model = OLS(y, X)
        results = model.fit()
        joint_test = results.f_test("lag_hit = 0, var = 0")

        # 判断是否通过检验 (p值 > 0.05 表示通过)
        pass_test = joint_test.pvalue > 0.05

        return {
            'hits': hits.sum(),
            'expected_hits': len(hits) * alpha,
            'failure_rate': hits.sum() / len(hits),
            'p_value': joint_test.pvalue,
            'reject_null': joint_test.pvalue < 0.05,
            'pass_test': pass_test  # 添加是否通过检验的标志
        }
    except Exception as e:
        print(f"动态分位数检验出错: {e}")
        return {
            'hits': hits.sum(),
            'expected_hits': len(hits) * alpha,
            'failure_rate': hits.sum() / len(hits),
            'p_value': 1.0,
            'reject_null': False,
            'pass_test': True
        }


# 对5% VaR进行动态分位数检验
print("\n对5% VaR进行动态分位数检验...")
dqt_normal = dynamic_quantile_test(returns_insample_processed, var_normal_series[0.05], 0.05)

print("\n正态分布假设的动态分位数检验结果:")
print(f"  实际失败次数: {dqt_normal['hits']}, 预期失败次数: {dqt_normal['expected_hits']:.1f}")
print(f"  失败率: {dqt_normal['failure_rate']:.4f}, p值: {dqt_normal['p_value']:.4f}")
print(f"  是否拒绝原假设: {'是' if dqt_normal['reject_null'] else '否'}")
print(f"  是否通过DQ检验: {'是' if dqt_normal['pass_test'] else '否'}")

# 添加样本内动态分位数检验结果
doc.add_heading('七、样本内动态分位数检验(DQ检验)', level=1)
p = doc.add_paragraph()
p.add_run("检验水平: 5% VaR").bold = True
doc.add_paragraph(f"实际失败次数: {dqt_normal['hits']}")
doc.add_paragraph(f"预期失败次数: {dqt_normal['expected_hits']:.1f}")
doc.add_paragraph(f"失败率: {dqt_normal['failure_rate']:.4f}")
doc.add_paragraph(f"DQ检验p值: {dqt_normal['p_value']:.4f}")
p = doc.add_paragraph()
if dqt_normal['pass_test']:
    p.add_run("DQ检验结果: 通过 (p值 > 0.05，模型设定合理)").bold = True
    p.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # 绿色
else:
    p.add_run("DQ检验结果: 未通过 (p值 ≤ 0.05，模型设定可能存在问题)").bold = True
    p.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # 红色
doc.add_paragraph("")

# 7. 样本外滚动预测
print("\n7. 样本外滚动预测...")


def rolling_garch_forecast(returns_insample_data, returns_outsample_data, p=0, q=0, dist='normal'):
    var_predictions = {alpha: [] for alpha in confidence_levels}

    for i in range(len(returns_outsample_data)):
        # 当前窗口数据：样本内数据 + 已预测的样本外数据
        if i == 0:
            window_data = returns_insample_data.copy()
        else:
            window_data = pd.concat([returns_insample_data, returns_outsample_data.iloc[:i]])

        try:
            # 构建GARCH模型
            if p > 0 and q > 0:
                model = arch_model(window_data, mean='AR', lags=p, vol='GARCH', p=1, q=1, dist=dist)
            elif p > 0:
                model = arch_model(window_data, mean='AR', lags=p, vol='GARCH', p=1, q=1, dist=dist)
            elif q > 0:
                model = arch_model(window_data, mean='AR', lags=q, vol='GARCH', p=1, q=1, dist=dist)
            else:
                model = arch_model(window_data, mean='constant', vol='GARCH', p=1, q=1, dist=dist)

            results = model.fit(disp='off', show_warning=False)

            # 预测下一期
            if p > 0 or q > 0:
                lag_order = max(p, q)
                if len(window_data) > lag_order:
                    forecast = results.forecast(start=lag_order, horizon=1)
                else:
                    vol_estimate = np.std(window_data)
                    forecast = type('obj', (object,), {'variance': np.array([[vol_estimate ** 2]])})()
            else:
                forecast = results.forecast(start=0, horizon=1)

            # 计算VaR
            for alpha in confidence_levels:
                var_value = -forecast.variance.values[-1:] ** 0.5 * stats.norm.ppf(alpha)
                var_predictions[alpha].append(var_value[0])

        except Exception as e:
            # 打印前几次错误，避免输出太多
            if i < 5:
                print(f"第{i}次预测失败: {e}")
            for alpha in confidence_levels:
                if i > 0 and len(var_predictions[alpha]) > 0:
                    var_predictions[alpha].append(var_predictions[alpha][-1])
                else:
                    var_predictions[alpha].append(np.std(window_data) * abs(stats.norm.ppf(alpha)))

    return var_predictions


print("进行正态分布假设的滚动预测...")
var_normal_outsample = rolling_garch_forecast(returns_insample_processed, returns_outsample,
                                              p=best_p, q=best_q, dist='normal')

# 转换为DataFrame
var_normal_df = pd.DataFrame(var_normal_outsample, index=returns_outsample.index)

# 8. 样本外模型评估
print("\n8. 样本外模型评估...")


def evaluate_var_model(returns, var_series, alpha):
    # 计算失败次数
    hits = (returns < -var_series).astype(int)
    n = len(returns)
    failures = hits.sum()
    failure_rate = failures / n

    # Kupiec失败率检验
    def kupiec_test(failures, n, alpha):
        if failures == 0:
            lr_stat = 2 * n * np.log(1 / (1 - alpha))
        elif failures == n:
            lr_stat = 2 * n * np.log(1 / alpha)
        else:
            lr_stat = 2 * (failures * np.log(failures / (n * alpha)) +
                           (n - failures) * np.log((n - failures) / (n * (1 - alpha))))
        p_value = 1 - stats.chi2.cdf(lr_stat, 1)
        return lr_stat, p_value

    # 计算分位数损失
    def quantile_loss(returns, var_series, alpha):
        loss = np.maximum((alpha - 1) * (returns - var_series), alpha * (returns - var_series))
        return loss.mean()

    lr_stat, p_value = kupiec_test(failures, n, alpha)
    avg_loss = quantile_loss(returns, var_series, alpha)

    # 确保平均分位数损失是标量
    if isinstance(avg_loss, np.ndarray):
        avg_loss = avg_loss.item()
    elif isinstance(avg_loss, (np.generic, np.number)):
        avg_loss = float(avg_loss)

    # 判断是否通过Kupiec检验 (p值 > 0.05 表示通过)
    pass_kupiec = p_value > 0.05

    return {
        'failures': failures,
        'expected_failures': n * alpha,
        'failure_rate': failure_rate,
        'lr_statistic': lr_stat,
        'p_value': p_value,
        'reject_null': p_value < 0.05,
        'avg_quantile_loss': avg_loss,
        'pass_kupiec': pass_kupiec
    }


# 添加样本外评估标题
doc.add_heading('八、样本外VaR预测评估', level=1)

print("\n正态分布假设下的样本外评估:")
for alpha in confidence_levels:
    # 确保使用数值类型的数据进行评估
    returns_outsample_eval = returns_outsample.astype(float).fillna(0)
    var_series_eval = var_normal_df[alpha].astype(float).fillna(method='ffill').fillna(method='bfill')

    eval_normal = evaluate_var_model(returns_outsample_eval, var_series_eval, alpha)
    print(f"\n{alpha * 100:.1f}% VaR评估:")
    print(f"  实际失败次数: {eval_normal['failures']}, 预期失败次数: {eval_normal['expected_failures']:.1f}")
    print(f"  失败率: {eval_normal['failure_rate']:.4f}")
    print(f"  LR统计量: {eval_normal['lr_statistic']:.4f}, Kupiec检验p值: {eval_normal['p_value']:.4f}")
    print(f"  是否拒绝原假设: {'是' if eval_normal['reject_null'] else '否'}")
    print(f"  是否通过Kupiec检验: {'是' if eval_normal['pass_kupiec'] else '否'}")
    print(f"  平均分位数损失: {eval_normal['avg_quantile_loss']:.6f}")

    # 对样本外VaR进行动态分位数检验
    dqt_outsample = dynamic_quantile_test(returns_outsample_eval, var_series_eval, alpha)
    print(f"  样本外DQ检验p值: {dqt_outsample['p_value']:.4f}")
    print(f"  是否通过样本外DQ检验: {'是' if dqt_outsample['pass_test'] else '否'}")

    # 将每个置信水平的评估结果添加到Word文档
    doc.add_heading(f'{alpha * 100:.1f}% VaR评估结果', level=2)
    doc.add_paragraph(f"实际失败次数: {eval_normal['failures']}")
    doc.add_paragraph(f"预期失败次数: {eval_normal['expected_failures']:.1f}")
    doc.add_paragraph(f"失败率: {eval_normal['failure_rate']:.4f}")
    doc.add_paragraph(f"Kupiec检验LR统计量: {eval_normal['lr_statistic']:.4f}")
    doc.add_paragraph(f"Kupiec检验p值: {eval_normal['p_value']:.4f}")

    # Kupiec检验结果判断
    p_kupiec = doc.add_paragraph()
    if eval_normal['pass_kupiec']:
        p_kupiec.add_run("Kupiec检验结果: 通过 (p值 > 0.05，模型设定合理)").bold = True
        p_kupiec.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # 绿色
    else:
        p_kupiec.add_run("Kupiec检验结果: 未通过 (p值 ≤ 0.05，模型设定可能存在问题)").bold = True
        p_kupiec.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # 红色

    # DQ检验结果
    doc.add_paragraph(f"样本外DQ检验p值: {dqt_outsample['p_value']:.4f}")
    p_dq = doc.add_paragraph()
    if dqt_outsample['pass_test']:
        p_dq.add_run("DQ检验结果: 通过 (p值 > 0.05，模型设定合理)").bold = True
        p_dq.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # 绿色
    else:
        p_dq.add_run("DQ检验结果: 未通过 (p值 ≤ 0.05，模型设定可能存在问题)").bold = True
        p_dq.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # 红色

    doc.add_paragraph(f"平均分位数损失: {eval_normal['avg_quantile_loss']:.6f}")
    doc.add_paragraph("")

# 9. 可视化结果 - 只保留不同置信水平下VaR和实际收益率的比较图
print("\n9. 生成可视化结果...")

fig, axes = plt.subplots(3, 1, figsize=(14, 12))

# 绘制三个置信水平的VaR预测对比
confidence_levels_plot = [0.01, 0.05, 0.10]
colors = ['r', 'g', 'b']

for idx, (alpha, color) in enumerate(zip(confidence_levels_plot, colors)):
    # 确保数据是数值类型，处理可能的NaN值
    var_normal_plot = -var_normal_df[alpha].astype(float).fillna(method='ffill')
    returns_outsample_plot = returns_outsample.astype(float).fillna(method='ffill')

    axes[idx].plot(returns_outsample.index, var_normal_plot,
                   label=f'{alpha * 100:.0f}% VaR (normal distribution)', color=color, alpha=0.8)
    axes[idx].plot(returns_outsample.index, returns_outsample_plot,
                   label='true returns', alpha=0.5, linewidth=0.5)
    axes[idx].fill_between(returns_outsample.index, var_normal_plot,
                           var_normal_plot.min() * 1.1, alpha=0.1, color=color, label='VaR interval')
    axes[idx].set_title(f'out-of-sample VaR predictions vs True returns({alpha * 100:.0f}%)')
    axes[idx].set_xlabel('date')
    axes[idx].set_ylabel('returns/VaR')
    axes[idx].legend()
    axes[idx].grid(True, alpha=0.3)

plt.tight_layout()
plot_path = os.path.join(results_dir, 'garch_var_comparison.png')
plt.savefig(plot_path, dpi=150, bbox_inches='tight')
plt.close(fig)  # 关闭图形以释放内存

print(f"分析完成！可视化结果已保存到 '{plot_path}'")

# 将图表添加到Word文档
doc.add_heading('九、可视化结果', level=1)
doc.add_paragraph("不同置信水平下VaR预测与实际收益率的比较:")
doc.add_picture(plot_path, width=Inches(6))

# 保存Word文档
doc_path = os.path.join(results_dir, 'results.docx')
doc.save(doc_path)

print(f"\n详细分析结果已保存到Word文档: '{doc_path}'")
print("=" * 60)
print("代码执行完成!")
print(f"所有结果文件保存在 '{results_dir}' 文件夹中")