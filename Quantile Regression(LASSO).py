import numpy as np
import pandas as pd
from scipy import stats
from statsmodels.tsa.stattools import adfuller
from statsmodels.stats.outliers_influence import variance_inflation_factor
from statsmodels.regression.linear_model import OLS
import statsmodels.api as sm
from scipy.optimize import minimize
import warnings
from docx import Document
from docx.shared import Inches
import os

warnings.filterwarnings('ignore')


class QuantileRegressionVaR:
    def __init__(self, file_path, tau_values=[0.01, 0.05, 0.10]):
        """
        初始化分位数回归VaR模型
        """
        # 读取数据
        try:
            self.data = pd.read_csv(file_path)
        except:
            # 尝试不同的分隔符
            self.data = pd.read_csv(file_path, sep='\t')

        print(f"数据形状: {self.data.shape}")
        print("数据列名:", self.data.columns.tolist())

        # 根据数据实际情况处理
        if len(self.data.columns) == 4:
            self.data.columns = ['date', 'log_ret', 'rv5', 'bv']
            self.data['date'] = pd.to_datetime(self.data['date'])
        elif len(self.data.columns) == 3:
            self.data.columns = ['log_ret', 'rv5', 'bv']
            # 如果没有日期列，创建一个
            self.data['date'] = pd.date_range(start='2000-01-04', periods=len(self.data), freq='B')  # 工作日频率
        else:
            # 假设第一列是日期，后面是数据
            self.data.columns = ['date', 'log_ret', 'rv5', 'bv'] + list(self.data.columns[4:])
            self.data['date'] = pd.to_datetime(self.data['date'])

        self.data = self.data.set_index('date')

        self.tau_values = tau_values
        self.rv_col = 'rv5'
        self.bv_col = 'bv'
        self.ret_col = 'log_ret'

        # 特征工程
        self._feature_engineering()

        # 划分样本
        self._split_sample()

        # 模型参数 - 调整正则化参数，避免过度压缩
        self.models = {tau: None for tau in tau_values}
        self.params = {tau: None for tau in tau_values}
        self.lambdas = {0.01: 0.0001, 0.05: 0.00005, 0.10: 0.00001}

        # 存储结果
        self.results = {}

    def _feature_engineering(self):
        """特征工程"""
        print("\n进行特征工程...")

        # 创建极端波动指标
        self.data['rv5_extreme'] = (
                self.data[self.rv_col] >
                self.data[self.rv_col].rolling(22, min_periods=1).quantile(0.9)
        ).astype(int)

        # 开根号转换
        self.data['rv_sqrt'] = np.sqrt(self.data[self.rv_col])
        self.data['bv_sqrt'] = np.sqrt(self.data[self.bv_col])

        # 计算rv5-bv差值并开根号
        self.data['rv5-bv_sqrt'] = np.sqrt(np.maximum(
            self.data[self.rv_col] - self.data[self.bv_col], 1e-10))

        # 极端波动指标开根号
        self.data['rv5_extreme_sqrt'] = np.sqrt(self.data['rv5_extreme'])

        # 5日移动平均并开根号
        self.data['rv5_ma5'] = self.data[self.rv_col].rolling(5, min_periods=1).mean()
        self.data['bv_ma5'] = self.data[self.bv_col].rolling(5, min_periods=1).mean()
        self.data['rv5_ma5_sqrt'] = np.sqrt(self.data['rv5_ma5'])
        self.data['bv_ma5_sqrt'] = np.sqrt(self.data['bv_ma5'])

        # 创建滞后一期的自变量
        lag_features = [
            'rv_sqrt', 'bv_sqrt', 'rv5-bv_sqrt',
            'rv5_extreme_sqrt', 'rv5_ma5_sqrt', 'bv_ma5_sqrt'
        ]

        for feature in lag_features:
            self.data[f'{feature}_lag1'] = self.data[feature].shift(1)

        # 定义自变量
        self.X_cols = [
            'rv_sqrt_lag1', 'bv_sqrt_lag1', 'rv5-bv_sqrt_lag1',
            'rv5_extreme_sqrt_lag1', 'rv5_ma5_sqrt_lag1', 'bv_ma5_sqrt_lag1'
        ]

        # 删除包含NaN的行
        original_len = len(self.data)
        self.data = self.data.dropna()
        new_len = len(self.data)
        print(f"删除NaN行: {original_len} -> {new_len} (删除{original_len - new_len}行)")

    def _split_sample(self):
        """按前80%和后20%划分样本"""
        n = len(self.data)
        split_idx = int(n * 0.8)

        self.train_data = self.data.iloc[:split_idx].copy()
        self.test_data = self.data.iloc[split_idx:].copy()

        print(f"\n样本划分:")
        print(f"总样本大小: {n}")
        print(f"训练样本大小: {len(self.train_data)} ({len(self.train_data) / n * 100:.1f}%)")
        print(f"测试样本大小: {len(self.test_data)} ({len(self.test_data) / n * 100:.1f}%)")
        print(f"训练期间: {self.train_data.index[0]} 到 {self.train_data.index[-1]}")
        print(f"测试期间: {self.test_data.index[0]} 到 {self.test_data.index[-1]}")

    def diagnostic_tests(self):
        """单位根检验和多重共线性检验"""
        print("\n" + "=" * 70)
        print("诊断检验")
        print("=" * 70)

        # 单位根检验（ADF检验）
        print("\n1. 单位根检验 (ADF检验, 5%显著性水平):")
        print("-" * 60)
        adf_results = []
        for col in self.X_cols:
            try:
                result = adfuller(self.train_data[col].dropna(), autolag='AIC')
                is_stationary = result[1] < 0.05
                adf_results.append({
                    '变量': col,
                    'ADF统计量': result[0],
                    'p值': result[1],
                    '是否平稳': '是' if is_stationary else '否'
                })
                star = "***" if is_stationary else ""
                print(
                    f"{col:25s}: ADF统计量={result[0]:8.4f}, p值={result[1]:.6f} {'(平稳)' if is_stationary else '(非平稳)'} {star}")
            except Exception as e:
                print(f"{col:25s}: 检验失败 - {e}")

        # 多重共线性检验（VIF）
        print("\n\n2. 多重共线性检验 (VIF):")
        print("-" * 60)
        X = self.train_data[self.X_cols]

        # 添加常数项
        X_with_const = sm.add_constant(X)

        vif_data = pd.DataFrame()
        vif_data["变量"] = ['常数项'] + self.X_cols
        vif_data["VIF"] = [variance_inflation_factor(X_with_const.values, i)
                           for i in range(X_with_const.shape[1])]

        # 添加诊断标记
        vif_data["诊断"] = vif_data["VIF"].apply(
            lambda x: "严重共线性(VIF>10)" if x > 10 else ("中度共线性(VIF>5)" if x > 5 else "正常")
        )

        print(vif_data.to_string(index=False))

        return pd.DataFrame(adf_results), vif_data

    def quantile_loss(self, y_true, y_pred, tau):
        """分位数损失函数"""
        errors = y_true - y_pred
        return np.where(errors < 0, (tau - 1) * errors, tau * errors)

    def objective_function(self, params, X, y, tau, lambda_reg):
        """目标函数：分位数损失 + LASSO正则化"""
        y_pred = X @ params
        loss = np.mean(self.quantile_loss(y, y_pred, tau))
        # 不对截距项进行正则化
        lasso_penalty = lambda_reg * np.sum(np.abs(params[1:]))
        return loss + lasso_penalty

    def train_models(self):
        """训练所有分位数模型（样本内）"""
        print("\n" + "=" * 70)
        print("样本内分位数回归模型训练")
        print("=" * 70)

        # 准备训练数据
        X_train = self.train_data[self.X_cols].values
        y_train = self.train_data[self.ret_col].values

        # 对样本内数据进行标准化
        self.X_train_mean = X_train.mean(axis=0)
        self.X_train_std = X_train.std(axis=0)
        self.X_train_std[self.X_train_std == 0] = 1.0  # 避免除零
        X_train_std = (X_train - self.X_train_mean) / self.X_train_std

        # 添加常数项
        X_train_with_const = np.column_stack([np.ones(X_train_std.shape[0]), X_train_std])

        for tau in self.tau_values:
            print(f"\n训练 {tau * 100:.0f}% 分位数模型 (τ={tau}, λ={self.lambdas[tau]}):")
            print("-" * 50)

            # 初始参数
            n_features = X_train_with_const.shape[1]
            init_params = np.zeros(n_features)

            # 优化
            bounds = [(None, None)] * n_features
            result = minimize(
                self.objective_function,
                init_params,
                args=(X_train_with_const, y_train, tau, self.lambdas[tau]),
                method='L-BFGS-B',
                bounds=bounds,
                options={'maxiter': 1000, 'ftol': 1e-8}
            )

            if result.success:
                self.params[tau] = result.x

                # 打印参数
                param_names = ['截距'] + self.X_cols
                print(f"{'参数':25s} {'值':>12s} ")
                print("-" * 50)
                for i, (name, value) in enumerate(zip(param_names, result.x)):
                    print(f"{name:25s} {value:12.6f}")

                # 统计非零系数
                non_zero = np.sum(np.abs(result.x[1:]) > 1e-6)
                print(f"\n非零系数数量: {non_zero}/{len(result.x) - 1}")
            else:
                print(f"优化失败: {result.message}")

    def dynamic_quantile_test(self, returns, var_predictions, tau, lags=5):
        """
        动态分位数检验（DQ检验）

        原假设：模型设定正确，能够充分捕捉数据动态
        备择假设：模型存在设定误差

        检验通过条件：p值 > 0.05（不拒绝原假设）
        """
        # 确保长度一致
        min_len = min(len(returns), len(var_predictions))
        returns = returns[:min_len]
        var_predictions = var_predictions[:min_len]

        # 计算失败指示函数
        hits = (returns < var_predictions).astype(int)
        hit_rate = np.mean(hits)
        expected_hit_rate = tau

        # 构建回归矩阵
        n = len(returns)
        if n <= lags + 1:
            print(f"样本量{n}不足，无法进行{lags}阶DQ检验")
            return None

        X = np.ones((n - lags, 2 + lags))

        for i in range(lags, n):
            X[i - lags, 1] = var_predictions[i]
            for j in range(lags):
                X[i - lags, 2 + j] = hits[i - j - 1]

        y = hits[lags:]

        # OLS回归
        try:
            model = OLS(y, X).fit()

            # 计算DQ统计量
            dq_stat = model.nobs * model.rsquared
            p_value = 1 - stats.chi2.cdf(dq_stat, df=lags + 1)

            # 检验结果
            test_passed = p_value > 0.05  # 不拒绝原假设

            return {
                'DQ_statistic': dq_stat,
                'p_value': p_value,
                'test_passed': test_passed,
                'hit_rate': hit_rate,
                'expected_hit_rate': expected_hit_rate,
                'hits': hits,
                'nobs': model.nobs,
                'rsquared': model.rsquared
            }
        except Exception as e:
            print(f"DQ检验失败: {e}")
            return None

    def kupiec_test(self, returns, var_predictions, tau):
        """
        Kupiec失败率检验

        原假设：VaR预测准确，失败率等于期望水平
        备择假设：VaR预测不准确，失败率不等于期望水平

        检验通过条件：p值 > 0.05（不拒绝原假设）
        """
        # 确保长度一致
        min_len = min(len(returns), len(var_predictions))
        returns = returns[:min_len]
        var_predictions = var_predictions[:min_len]

        # 计算失败次数
        hits = (returns < var_predictions).astype(int)
        n = len(hits)
        x = np.sum(hits)
        failure_rate = x / n if n > 0 else 0

        # 计算似然比统计量
        if x == 0:
            lr = -2 * n * np.log(1 - tau)
        elif x == n:
            lr = -2 * n * np.log(tau)
        else:
            p_hat = x / n
            lr = -2 * (np.log((1 - tau) ** (n - x) * tau ** x) -
                       np.log((1 - p_hat) ** (n - x) * p_hat ** x))

        p_value = 1 - stats.chi2.cdf(lr, df=1)

        # 检验结果
        test_passed = p_value > 0.05  # 不拒绝原假设

        return {
            'failures': x,
            'total': n,
            'failure_rate': failure_rate,
            'expected_rate': tau,
            'LR_statistic': lr,
            'p_value': p_value,
            'test_passed': test_passed,
            'hits': hits
        }

    def rolling_forecast_with_standardization(self):
        """
        样本外滚动预测（每个窗口独立标准化）

        在每个滚动窗口内，使用该窗口的数据计算均值和标准差进行标准化
        """
        print("\n" + "=" * 70)
        print("样本外滚动预测（滚动标准化）")
        print("=" * 70)

        test_size = len(self.test_data)

        print(f"测试集大小: {test_size}")
        print(f"总预测次数: {test_size}")
        print(f"滚动标准化: 每个窗口独立计算标准化参数")

        # 初始化存储
        forecasts = {tau: np.zeros(test_size) for tau in self.tau_values}
        losses = {tau: [] for tau in self.tau_values}
        all_window_params = {tau: [] for tau in self.tau_values}

        # 滚动预测
        for i in range(test_size):
            # 创建滚动窗口（训练集 + 已预测的测试集部分）
            if i == 0:
                window_data = self.train_data.copy()
            else:
                window_data = pd.concat([
                    self.train_data,
                    self.test_data.iloc[:i]
                ])

            # 显示进度
            if (i + 1) % 100 == 0 or i == test_size - 1:
                print(f"进度: {i + 1}/{test_size} ({((i + 1) / test_size * 100):.1f}%)")

            # 对当前窗口进行标准化
            X_window = window_data[self.X_cols].values
            y_window = window_data[self.ret_col].values

            # 计算当前窗口的标准化参数
            X_mean = X_window.mean(axis=0)
            X_std = X_window.std(axis=0)
            X_std[X_std == 0] = 1.0  # 避免除零

            # 标准化当前窗口数据
            X_window_std = (X_window - X_mean) / X_std

            # 添加常数项
            X_window_with_const = np.column_stack([np.ones(X_window_std.shape[0]), X_window_std])

            for tau in self.tau_values:
                try:
                    # 训练模型
                    n_features = X_window_with_const.shape[1]
                    init_params = np.zeros(n_features)

                    # 使用较小的正则化参数
                    lambda_reg = self.lambdas[tau] * 0.1  # 进一步减小，避免系数压缩

                    bounds = [(None, None)] * n_features
                    result = minimize(
                        lambda params: np.mean(self.quantile_loss(y_window, X_window_with_const @ params, tau)) +
                                       lambda_reg * np.sum(np.abs(params[1:])),
                        init_params,
                        method='L-BFGS-B',
                        bounds=bounds,
                        options={'maxiter': 500, 'ftol': 1e-6}
                    )

                    if result.success:
                        # 保存窗口参数
                        all_window_params[tau].append(result.x.copy())

                        # 预测下一个观测值
                        X_test = self.test_data.iloc[i:i + 1][self.X_cols].values
                        # 使用当前窗口的标准化参数
                        X_test_std = (X_test - X_mean) / X_std
                        X_test_with_const = np.column_stack([np.ones(X_test_std.shape[0]), X_test_std])

                        forecast = X_test_with_const @ result.x
                        forecasts[tau][i] = forecast[0]

                        # 计算损失
                        actual = self.test_data.iloc[i][self.ret_col]
                        loss = self.quantile_loss(actual, forecast, tau)
                        losses[tau].append(loss)
                    else:
                        # 如果优化失败，使用简单分位数
                        forecasts[tau][i] = np.percentile(y_window, tau * 100)
                        losses[tau].append(0)
                        all_window_params[tau].append(None)

                except Exception as e:
                    print(f"  第{i + 1}次预测错误 ({tau}分位数): {e}")
                    # 使用简单分位数作为后备
                    forecasts[tau][i] = np.percentile(y_window, tau * 100)
                    losses[tau].append(0)
                    all_window_params[tau].append(None)

        return forecasts, losses, all_window_params

    def evaluate_models(self):
        """全面评估模型性能"""
        print("\n" + "=" * 70)
        print("模型综合评估")
        print("=" * 70)

        # 1. 样本内预测和检验
        print("\n1. 样本内模型性能:")
        print("-" * 50)

        # 准备样本内数据
        X_train = self.train_data[self.X_cols].values
        X_train_std = (X_train - self.X_train_mean) / self.X_train_std
        X_train_with_const = np.column_stack([np.ones(X_train_std.shape[0]), X_train_std])
        y_train = self.train_data[self.ret_col].values

        sample_in_results = {}

        for tau in self.tau_values:
            print(f"\n{tau * 100:.0f}% VaR 样本内:")
            print("-" * 30)

            if self.params[tau] is not None:
                # 样本内预测
                var_train = X_train_with_const @ self.params[tau]

                # 计算样本内失败率
                hits_train = (y_train < var_train).astype(int)
                hit_rate_train = np.mean(hits_train)

                # DQ检验
                dq_result_in = self.dynamic_quantile_test(y_train, var_train, tau)

                if dq_result_in:
                    dq_status = "通过" if dq_result_in['test_passed'] else "未通过"
                    print(f"  DQ检验: 统计量={dq_result_in['DQ_statistic']:.4f}, "
                          f"p值={dq_result_in['p_value']:.4f}, 结果={dq_status}")
                    print(f"  失败率: 实际={hit_rate_train:.4f}, 期望={tau:.4f}")

                sample_in_results[tau] = {
                    'var_predictions': var_train,
                    'actual_returns': y_train,
                    'hit_rate': hit_rate_train,
                    'dq_test': dq_result_in
                }

        # 2. 样本外滚动预测
        print("\n\n2. 样本外滚动预测性能:")
        print("-" * 50)

        forecasts, losses, window_params = self.rolling_forecast_with_standardization()

        sample_out_results = {}

        for tau in self.tau_values:
            print(f"\n{tau * 100:.0f}% VaR 样本外:")
            print("-" * 30)

            var_pred = forecasts[tau]
            actual_test = self.test_data[self.ret_col].values[:len(var_pred)]

            # Kupiec检验
            kupiec_result = self.kupiec_test(actual_test, var_pred, tau)

            # DQ检验
            dq_result_out = self.dynamic_quantile_test(actual_test, var_pred, tau)

            if kupiec_result and dq_result_out:
                kupiec_status = "通过" if kupiec_result['test_passed'] else "未通过"
                dq_status = "通过" if dq_result_out['test_passed'] else "未通过"

                print(f"  Kupiec检验: 失败次数={kupiec_result['failures']}/{kupiec_result['total']}, "
                      f"失败率={kupiec_result['failure_rate']:.4f}, 期望={tau:.4f}")
                print(f"              LR统计量={kupiec_result['LR_statistic']:.4f}, "
                      f"p值={kupiec_result['p_value']:.4f}, 结果={kupiec_status}")
                print(f"  DQ检验: 统计量={dq_result_out['DQ_statistic']:.4f}, "
                      f"p值={dq_result_out['p_value']:.4f}, 结果={dq_status}")

                # 平均损失
                avg_loss = np.mean(losses[tau])
                print(f"  平均分位数损失: {avg_loss:.6f}")

                # VaR统计
                print(f"  VaR统计: 均值={np.mean(var_pred):.6f}, 标准差={np.std(var_pred):.6f}")

            sample_out_results[tau] = {
                'var_predictions': var_pred,
                'actual_returns': actual_test,
                'kupiec_test': kupiec_result,
                'dq_test': dq_result_out,
                'avg_loss': np.mean(losses[tau]) if losses[tau] else 0,
                'window_params': window_params[tau]
            }

        # 3. 总结对比
        print("\n\n3. 样本内 vs 样本外 性能对比:")
        print("-" * 50)
        print(f"{'分位数':<8} {'样本':<8} {'失败率':<12} {'DQ检验':<10} {'Kupiec检验':<12} {'平均损失':<12}")
        print("-" * 70)

        # 收集数据用于Word输出
        performance_data = []

        for tau in self.tau_values:
            # 样本内
            if tau in sample_in_results and sample_in_results[tau]['dq_test']:
                in_dq_status = "通过" if sample_in_results[tau]['dq_test']['test_passed'] else "未通过"
                print(f"{tau * 100:>4.0f}%   {'样本内':<8} {sample_in_results[tau]['hit_rate']:>10.4f}  "
                      f"{in_dq_status:>10} {'-':>12} {'-':>12}")

                # 收集数据
                performance_data.append({
                    '分位数': f"{tau * 100:.0f}%",
                    '样本': '样本内',
                    '失败率': sample_in_results[tau]['hit_rate'],
                    'DQ检验': in_dq_status,
                    'Kupiec检验': '-',
                    '平均损失': '-'
                })

            # 样本外
            if tau in sample_out_results and sample_out_results[tau]['kupiec_test'] and sample_out_results[tau][
                'dq_test']:
                out_dq_status = "通过" if sample_out_results[tau]['dq_test']['test_passed'] else "未通过"
                out_kupiec_status = "通过" if sample_out_results[tau]['kupiec_test']['test_passed'] else "未通过"
                print(
                    f"{tau * 100:>4.0f}%   {'样本外':<8} {sample_out_results[tau]['kupiec_test']['failure_rate']:>10.4f}  "
                    f"{out_dq_status:>10} {out_kupiec_status:>12} {sample_out_results[tau]['avg_loss']:>12.6f}")

                # 收集数据
                performance_data.append({
                    '分位数': f"{tau * 100:.0f}%",
                    '样本': '样本外',
                    '失败率': sample_out_results[tau]['kupiec_test']['failure_rate'],
                    'DQ检验': out_dq_status,
                    'Kupiec检验': out_kupiec_status,
                    '平均损失': sample_out_results[tau]['avg_loss']
                })

        # 保存结果
        self.results = {
            'sample_in': sample_in_results,
            'sample_out': sample_out_results,
            'performance_data': performance_data
        }

        return self.results

    def save_results_to_word(self):
        """将结果保存到Word文档"""
        # 创建QR results文件夹
        if not os.path.exists('QR results'):
            os.makedirs('QR results')

        # 创建Word文档
        doc = Document()

        # 添加标题
        doc.add_heading('分位数回归VaR模型分析结果', 0)

        # 添加基本信息
        doc.add_heading('1. 模型基本信息', level=1)
        doc.add_paragraph(f'分析时间: {pd.Timestamp.now()}')
        doc.add_paragraph(f'分析的分位数水平: {[f"{tau * 100:.0f}%" for tau in self.tau_values]}')

        # 添加样本划分信息
        doc.add_heading('2. 样本划分', level=1)
        doc.add_paragraph(f'训练样本大小: {len(self.train_data)}')
        doc.add_paragraph(f'测试样本大小: {len(self.test_data)}')
        doc.add_paragraph(f'训练期间: {self.train_data.index[0]} 到 {self.train_data.index[-1]}')
        doc.add_paragraph(f'测试期间: {self.test_data.index[0]} 到 {self.test_data.index[-1]}')

        # 添加性能对比表格
        doc.add_heading('3. 样本内 vs 样本外 性能对比', level=1)

        # 创建表格
        table = doc.add_table(rows=1, cols=6)
        table.style = 'LightShading-Accent1'

        # 设置表头
        header_cells = table.rows[0].cells
        headers = ['分位数', '样本', '失败率', 'DQ检验', 'Kupiec检验', '平均损失']
        for i, header in enumerate(headers):
            header_cells[i].text = header

        # 填充数据
        if 'performance_data' in self.results:
            for data in self.results['performance_data']:
                row_cells = table.add_row().cells
                row_cells[0].text = data['分位数']
                row_cells[1].text = data['样本']
                row_cells[2].text = f"{data['失败率']:.4f}" if isinstance(data['失败率'], (int, float)) else data[
                    '失败率']
                row_cells[3].text = data['DQ检验']
                row_cells[4].text = data['Kupiec检验']
                row_cells[5].text = f"{data['平均损失']:.6f}" if isinstance(data['平均损失'], (int, float)) else data[
                    '平均损失']

        # 添加检验说明
        doc.add_heading('4. 检验说明', level=1)
        doc.add_paragraph('DQ检验（动态分位数检验）:')
        doc.add_paragraph('   - 原假设：模型设定正确，能够充分捕捉数据动态')
        doc.add_paragraph('   - 检验通过条件：p值 > 0.05（不拒绝原假设）')
        doc.add_paragraph('   - 如果被拒绝，说明模型存在设定误差')
        doc.add_paragraph('Kupiec检验（失败率检验）:')
        doc.add_paragraph('   - 原假设：VaR预测准确，失败率等于期望水平')
        doc.add_paragraph('   - 检验通过条件：p值 > 0.05（不拒绝原假设）')
        doc.add_paragraph('   - 如果被拒绝，说明VaR预测不准确')

        # 添加最终总结
        doc.add_heading('5. 最终分析总结', level=1)

        for tau in self.tau_values:
            doc.add_heading(f'{tau * 100:.0f}% VaR 模型总结', level=2)

            if tau in self.results['sample_in'] and self.results['sample_in'][tau]['dq_test']:
                in_dq = self.results['sample_in'][tau]['dq_test']
                in_dq_pass = "✓通过" if in_dq['test_passed'] else "✗未通过"
                doc.add_paragraph(f'样本内 DQ检验: {in_dq_pass} (p值={in_dq["p_value"]:.4f})')

            if tau in self.results['sample_out']:
                out_data = self.results['sample_out'][tau]
                if out_data['kupiec_test']:
                    kupiec = out_data['kupiec_test']
                    kupiec_pass = "✓通过" if kupiec['test_passed'] else "✗未通过"
                    doc.add_paragraph(f'样本外 Kupiec检验: {kupiec_pass} (p值={kupiec["p_value"]:.4f})')
                    doc.add_paragraph(f'  失败率: {kupiec["failure_rate"]:.4f} (期望: {tau:.4f})')

                if out_data['dq_test']:
                    dq = out_data['dq_test']
                    dq_pass = "✓通过" if dq['test_passed'] else "✗未通过"
                    doc.add_paragraph(f'样本外 DQ检验: {dq_pass} (p值={dq["p_value"]:.4f})')

                doc.add_paragraph(f'样本外平均损失: {out_data["avg_loss"]:.6f}')

        # 保存文档
        doc_path = os.path.join('QR results', 'results.docx')
        doc.save(doc_path)
        print(f"\n结果已保存到: {doc_path}")

        return doc_path

    def run_complete_analysis(self):
        """运行完整的分析流程"""
        print("=" * 70)
        print("分位数回归VaR模型完整分析")
        print("=" * 70)

        # 1. 诊断检验
        adf_results, vif_results = self.diagnostic_tests()

        # 2. 训练样本内模型
        self.train_models()

        # 3. 综合评估
        results = self.evaluate_models()

        # 4. 保存结果到Word文档
        self.save_results_to_word()

        # 5. 最终总结
        print("\n" + "=" * 70)
        print("最终分析总结")
        print("=" * 70)

        print("\n\n各分位数水平模型表现总结:")
        print("=" * 70)

        for tau in self.tau_values:
            print(f"\n{tau * 100:.0f}% VaR 模型:")
            print("-" * 40)

            if tau in self.results['sample_in'] and self.results['sample_in'][tau]['dq_test']:
                in_dq = self.results['sample_in'][tau]['dq_test']
                in_dq_pass = "✓通过" if in_dq['test_passed'] else "✗未通过"
                print(f"样本内 DQ检验: {in_dq_pass} (p值={in_dq['p_value']:.4f})")

            if tau in self.results['sample_out']:
                out_data = self.results['sample_out'][tau]
                if out_data['kupiec_test']:
                    kupiec = out_data['kupiec_test']
                    kupiec_pass = "✓通过" if kupiec['test_passed'] else "✗未通过"
                    print(f"样本外 Kupiec检验: {kupiec_pass} (p值={kupiec['p_value']:.4f})")
                    print(f"  失败率: {kupiec['failure_rate']:.4f} (期望: {tau:.4f})")

                if out_data['dq_test']:
                    dq = out_data['dq_test']
                    dq_pass = "✓通过" if dq['test_passed'] else "✗未通过"
                    print(f"样本外 DQ检验: {dq_pass} (p值={dq['p_value']:.4f})")

                print(f"样本外平均损失: {out_data['avg_loss']:.6f}")

        print("\n" + "=" * 70)
        print("分析完成!")
        print("=" * 70)

        return self.results


# 主程序
if __name__ == "__main__":
    try:
        # 初始化模型
        model = QuantileRegressionVaR('spy_data.csv')

        # 运行完整分析
        results = model.run_complete_analysis()

    except FileNotFoundError:
        print("错误: 找不到文件 'spy_data.csv'")
        print("请确保文件在当前目录下，或提供完整的文件路径")
    except Exception as e:
        print(f"运行过程中出现错误: {e}")
        import traceback

        traceback.print_exc()

print(model.data[['rv5_extreme_sqrt','rv5_extreme_sqrt_lag1', 'rv5-bv_sqrt']].head())