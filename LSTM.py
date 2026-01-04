import numpy as np
import pandas as pd
import matplotlib as mlp
import matplotlib.pyplot as plt
import random
import statsmodels.api as sm
from scipy import stats  # 需要这个来计算p值
from sklearn.preprocessing import MinMaxScaler
from sklearn.metrics import mean_squared_error
import torch
import torch.nn as nn
from torch.autograd import Variable
from torch.utils.data import DataLoader, TensorDataset
import torch.optim as optim
import torch.utils.data as data
from docx import Document
from docx.shared import Inches
import os
import tempfile

# 读取CSV文件
df_spy_data = pd.read_csv('spy_data.csv')


# 在分割数据前进行特征处理
def transform_features(df):
    df_transformed = df.copy()
    # 对rv5和bv开根号
    df_transformed['rv5_sqrt'] = np.sqrt(df_transformed['rv5'])
    df_transformed['bv_sqrt'] = np.sqrt(df_transformed['bv'])
    df_transformed['log_ret%'] = df_transformed['log_ret'] * 100
    diff = df_transformed['rv5'] - df_transformed['bv']
    df_transformed['rv5-bv_sqrt'] = np.sqrt(np.maximum(diff, 0))
    # 删除原始列
    df_transformed = df_transformed.drop(['log_ret'], axis=1)
    # 重排列顺序
    df_transformed = df_transformed[['Unnamed: 0', 'log_ret%', 'rv5', 'bv', 'rv5_sqrt', 'bv_sqrt', 'rv5-bv_sqrt']]
    return df_transformed


# 应用特征变换
df_spy_data = transform_features(df_spy_data)

# 按时间顺序对训练集测试集进行分割
train_size = int(len(df_spy_data) * 0.6)
validation_size = int(len(df_spy_data) * 0.2)
test_size = len(df_spy_data) - train_size - validation_size
validation_end = train_size + validation_size
train, validation, test = df_spy_data[:train_size], df_spy_data[train_size:validation_end], df_spy_data[validation_end:]


def create_multivariate_dataset(dataset, window_size):
    """
    将时间序列进行滑窗
    参数:
        dataset: DataFrame，其中包含特征和标签，特征从索引1开始，第一列是标签
        window_size: 滑窗的窗口大小
    """
    X, y = [], []
    for i in range(len(dataset) - window_size):
        # 选取从第2列到最后一列的特征和标签
        feature_and_label = dataset.iloc[i:i + window_size, 1:].values  # 包括标签在内的特征
        # 下一个时间点的标签作为目标
        target = dataset.iloc[i + window_size, 1]  # 使用下一个时间点的标签作为预测目标
        X.append(feature_and_label)
        y.append(target)
    return torch.FloatTensor(np.array(X, dtype=np.float32)), torch.FloatTensor(np.array(y, dtype=np.float32))


X_train, y_train = create_multivariate_dataset(train, 21)
X_validation, y_validation = create_multivariate_dataset(validation, 21)
X_test, y_test = create_multivariate_dataset(test, 21)


def set_seed(seed_value=1):
    random.seed(seed_value)  # Python内置的随机库
    np.random.seed(seed_value)  # Numpy库
    torch.manual_seed(seed_value)  # 为CPU设置种子
    torch.cuda.manual_seed(seed_value)  # 为当前GPU设置种子
    torch.cuda.manual_seed_all(seed_value)  # 为所有GPU设置种子
    torch.backends.cudnn.deterministic = True  # 确保每次返回的卷积算法将是确定的
    torch.backends.cudnn.benchmark = False


# 模型架构
class Basic_LSTM(nn.Module):
    def __init__(self, input_dim, pred_len):
        super().__init__()
        self.lstm = nn.LSTM(input_size=input_dim, hidden_size=50, num_layers=1, batch_first=True)
        self.linear = nn.Linear(50, pred_len)
        self.pred_len = pred_len

    def forward(self, x):
        x, _ = self.lstm(x)
        x = self.linear(x[:, -1, :])
        if self.pred_len == 1:
            x = x.squeeze(-1)
        return x


# 定义分位数损失函数
class QuantileLoss(nn.Module):
    def __init__(self, alpha=0.01):
        super(QuantileLoss, self).__init__()
        self.alpha = alpha

    def forward(self, y_pred, y_true):
        errors = y_true - y_pred
        return torch.mean(torch.max((self.alpha - 1) * errors, self.alpha * errors))


# 定义动态分位数检验函数
def dq_test_with_pvalue(y, q_pred, tau, lags=4):
    """返回DQ检验的统计量和P值"""
    hit = (y <= q_pred).astype(float)
    # 构建回归矩阵
    z_t = np.column_stack([np.ones(len(hit) - lags)] +
                          [hit[i:-(lags - i)] for i in range(lags)])
    # DQ检验
    model = sm.OLS(hit[lags:] - tau, z_t).fit()
    # Wald检验
    n_coef = len(model.params)
    R = np.eye(n_coef)[1:, :]
    wald_result = model.wald_test(R, scalar=True)
    # 获取统计量和p值
    dq_statistic = wald_result.statistic
    dq_pvalue = wald_result.pvalue  # p值
    return dq_statistic, dq_pvalue


# 定义Kupiec检验函数
def kupiec_test_pytorch(y_true, y_pred, confidence_level=0.05):
    """
    Kupiec失败率检验（LR统计量）
    参数说明：
    y_true : torch.Tensor 或 np.array
        shape: (n_samples,)
    y_pred : torch.Tensor 或 np.array
        shape: (n_samples,)
   confidence_level : float
       预期失败率
    返回：
    dict: 包含检验结果的字典
        - 'n_samples': 总样本数
        - 'failures': 失败次数（实际收益率 < VaR）
        - 'expected_failures': 预期失败次数
        - 'failure_rate': 实际失败率
        - 'expected_rate': confidence_level
        - 'LR_statistic': LR统计量
        - 'p_value': p值
        - 'reject_null': 是否拒绝原假设（模型不准确）
    """
    # 确保为Tensor并移到CPU
    if isinstance(y_true, np.ndarray):
        y_true = torch.from_numpy(y_true)
    if isinstance(y_pred, np.ndarray):
        y_pred = torch.from_numpy(y_pred)

    # 计算失败次数（实际收益率 < VaR预测值）
    failures = torch.sum(y_true < y_pred).item()
    n_samples = y_true.shape[0]

    # 计算失败率
    failure_rate = failures / n_samples
    expected_rate = confidence_level

    # 计算预期失败次数
    expected_failures = n_samples * expected_rate

    # 计算LR统计量（似然比检验）
    # 原假设：实际失败率 = 预期失败率
    if failures == 0:
        LR = -2 * n_samples * torch.log(torch.tensor(1 - expected_rate))
    elif failures == n_samples:
        LR = -2 * n_samples * torch.log(torch.tensor(expected_rate))
    else:
        # LR = -2 * ln[(1-p)^(N-x) * p^x] + 2 * ln[(1-(x/N))^(N-x) * (x/N)^x]
        p = expected_rate
        x = failures
        N = n_samples

        # 使用对数避免数值下溢
        term1 = (N - x) * torch.log(torch.tensor(1 - p)) + x * torch.log(torch.tensor(p))
        term2 = (N - x) * torch.log(torch.tensor(1 - x / N)) + x * torch.log(torch.tensor(x / N))
        LR = -2 * (term1 - term2)

    # 计算p值（自由度为1的卡方分布）
    p_value = 1 - stats.chi2.cdf(LR.item(), df=1)

    # 判断结果（通常使用5%显著性水平）
    reject_null = p_value < 0.05

    # 返回结果字典
    return {
        'n_samples': n_samples,
        'failures': failures,
        'expected_failures': round(expected_failures, 2),
        'failure_rate': round(failure_rate, 4),
        'expected_rate': round(expected_rate, 4),
        'LR_statistic': round(LR.item(), 4),
        'p_value': round(p_value, 4),
        'reject_null': reject_null,
        'conclusion': '模型拒绝' if reject_null else '模型接受'
    }


# 设置参数
input_size = 6  # 输入特征的维度
hidden_size = 50  # LSTM隐藏状态的维度
num_layers = 1  # LSTM层的数量
output_size = 1  # 模型的输出维度
n_epochs = 200  # 迭代epoch
learning_rate = 0.001  # 学习率
alphas = [0.1, 0.05, 0.01]  # 定义要测试的多个置信水平


# 将整个训练、测试和结果保存过程封装到一个函数中
def train_and_evaluate_for_alpha(alpha, X_train, y_train, X_validation, y_validation, X_test, y_test,
                                 input_size=6, window_size=21):
    """
    针对给定的alpha值进行训练和评估

    参数:
        alpha: 置信水平
        X_train, y_train: 训练数据
        X_validation, y_validation: 验证数据
        X_test, y_test: 测试数据
        input_size: 输入特征维度
        window_size: 窗口大小

    返回:
        包含所有结果的字典
    """
    print(f"\n{'=' * 60}")
    print(f"开始训练和评估 alpha={alpha}")
    print(f"{'=' * 60}")

    # 设置随机种子以确保结果可重现
    set_seed(2)

    # 1. 初始化模型
    model = Basic_LSTM(input_dim=input_size, pred_len=1)
    optimizer = optim.Adam(model.parameters(), lr=0.001)
    loss_fn = QuantileLoss(alpha=alpha)

    # 2. 创建数据加载器
    loader = data.DataLoader(
        data.TensorDataset(X_train, y_train),
        shuffle=True,
        batch_size=8
    )

    # 3. 训练模型
    early_stopping_patience = 3
    early_stopping_counter = 0
    best_validation_QL = float('inf')
    train_losses = []
    val_losses = []

    for epoch in range(200):
        model.train()
        for X_batch, y_batch in loader:
            y_pred = model(X_batch)
            loss = loss_fn(y_pred, y_batch)
            optimizer.zero_grad()
            loss.backward()
            optimizer.step()

        # 每15个epoch记录一次损失
        if epoch % 15 == 0:
            model.eval()
            with torch.no_grad():
                y_train_pred = model(X_train)
                train_QL = loss_fn(y_train_pred, y_train)
                y_val_pred = model(X_validation)
                validation_QL = loss_fn(y_val_pred, y_validation)

            train_losses.append((epoch, train_QL.item()))
            val_losses.append((epoch, validation_QL.item()))

            print(f"Epoch {epoch}: Train QL {train_QL:.4f}, Validation QL {validation_QL:.4f}")

            # 早停检查
            if validation_QL < best_validation_QL:
                best_validation_QL = validation_QL
                early_stopping_counter = 0
            else:
                early_stopping_counter += 1
                if early_stopping_counter >= early_stopping_patience:
                    print(f"早停触发于 epoch {epoch}")
                    break

    # 4. 在测试集上进行预测
    model.eval()
    with torch.no_grad():
        y_pred_test = model(X_test)
        test_QL = loss_fn(y_pred_test, y_test)

    # 转换为numpy数组
    y_pred_np = y_pred_test.cpu().numpy().flatten()
    y_true_np = y_test.cpu().numpy().flatten()

    # 5. 计算DQ检验
    dq_stat, dq_p = dq_test_with_pvalue(y=y_true_np, q_pred=y_pred_np, tau=alpha, lags=4)

    # 6. 计算Kupiec检验
    kupiec_result = kupiec_test_pytorch(y_true_np, y_pred_np, confidence_level=alpha)

    # 7. 准备结果
    results = {
        'alpha': alpha,
        'model': model,
        'y_true': y_true_np,
        'y_pred': y_pred_np,
        'train_losses': train_losses,
        'val_losses': val_losses,
        'dq_statistic': dq_stat,
        'dq_pvalue': dq_p,
        'kupiec_result': kupiec_result,
        'test_QL': test_QL.item() if hasattr(test_QL, 'item') else test_QL
    }

    return results


# 创建保存结果的函数
def create_visualizations_and_save(results_dict, output_dir='LSTM results'):
    """
    为每个alpha值创建可视化并保存到Word文档

    参数:
        results_dict: 包含所有alpha结果的大字典
        output_dir: 输出目录
    """
    # 创建输出目录
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 创建Word文档
    doc = Document()
    doc.add_heading('LSTM VaR预测结果报告', 0)

    # 添加基本信息
    doc.add_heading('实验基本信息', level=1)

    # 修正：使用add_run而不是add_text
    p = doc.add_paragraph()
    p.add_run('模型类型: ').bold = True
    p.add_run('LSTM (长短期记忆网络)')

    p = doc.add_paragraph()
    p.add_run('数据: ').bold = True
    p.add_run('SPY股票数据')

    p = doc.add_paragraph()
    p.add_run('训练/验证/测试比例: ').bold = True
    p.add_run('60%/20%/20%')

    p = doc.add_paragraph()
    p.add_run('窗口大小: ').bold = True
    p.add_run('21天')

    doc.add_heading('各置信水平结果', level=1)

    # 为每个alpha创建章节
    for alpha, results in results_dict.items():
        doc.add_heading(f'置信水平 α = {alpha}', level=2)

        # 添加损失曲线
        doc.add_heading('训练过程损失曲线', level=3)

        # 创建损失曲线图
        plt.figure(figsize=(10, 5))
        epochs, train_loss = zip(*results['train_losses'])
        _, val_loss = zip(*results['val_losses'])

        plt.plot(epochs, train_loss, label=f'Train QL (α={alpha})', marker='o')
        plt.plot(epochs, val_loss, label=f'Validation QL (α={alpha})', marker='s')
        plt.xlabel('Epochs')
        plt.ylabel('Quantile Loss')
        plt.title(f'Training and Validation Loss (α={alpha})')
        plt.legend()
        plt.grid(True, alpha=0.3)

        # 保存图片
        loss_plot_path = os.path.join(output_dir, f'loss_alpha_{alpha}.png')
        plt.savefig(loss_plot_path, dpi=150, bbox_inches='tight')
        plt.close()

        # 添加到Word文档
        doc.add_picture(loss_plot_path, width=Inches(6))
        doc.add_paragraph(f'图: α={alpha} 的训练和验证损失曲线')

        # 添加测试集预测结果图
        doc.add_heading('测试集预测结果', level=3)

        plt.figure(figsize=(12, 6))
        time_indices = np.arange(len(results['y_true']))

        plt.plot(time_indices, results['y_true'], 'b-', label='True Values', alpha=0.7, linewidth=1)
        plt.plot(time_indices, results['y_pred'], 'r-', label=f'Predictions (α={alpha})', alpha=0.7, linewidth=1)

        plt.title(f'Test Set: True Values vs Predictions (α={alpha})', fontsize=14)
        plt.xlabel('Time Index', fontsize=12)
        plt.ylabel('log_ret%', fontsize=12)
        plt.legend(fontsize=10)
        plt.grid(True, alpha=0.3)

        # 保存图片
        test_plot_path = os.path.join(output_dir, f'test_alpha_{alpha}.png')
        plt.savefig(test_plot_path, dpi=150, bbox_inches='tight')
        plt.close()

        # 添加到Word文档
        doc.add_picture(test_plot_path, width=Inches(6))
        doc.add_paragraph(f'图: α={alpha} 的测试集预测结果')

        # 添加统计检验结果
        doc.add_heading('统计检验结果', level=3)

        # DQ检验结果
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Shading'

        # 设置表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '检验项目'
        hdr_cells[1].text = '结果'

        # DQ检验
        row_cells = table.rows[1].cells
        row_cells[0].text = 'DQ检验'
        dq_stat = results['dq_statistic']
        dq_p = results['dq_pvalue']
        dq_status = "通过 (p > 0.05)" if dq_p > 0.05 else "未通过 (p ≤ 0.05)"
        row_cells[1].text = f'统计量: {dq_stat:.4f}, p值: {dq_p:.4f}, 结果: {dq_status}'

        # Kupiec检验
        row_cells = table.rows[2].cells
        row_cells[0].text = 'Kupiec检验'
        kupiec = results['kupiec_result']
        row_cells[1].text = (f'失败率: {kupiec["failure_rate"]:.4f} (预期: {kupiec["expected_rate"]:.4f}), '
                             f'LR统计量: {kupiec["LR_statistic"]:.4f}, p值: {kupiec["p_value"]:.4f}, '
                             f'结论: {kupiec["conclusion"]}')

        # 添加模型性能总结
        doc.add_heading('模型性能总结', level=3)
        summary = doc.add_paragraph()
        summary.add_run('测试集损失: ').bold = True
        summary.add_run(f'{results["test_QL"]:.6f}\n')

        summary.add_run('DQ检验结果: ').bold = True
        summary.add_run(f'{dq_status}\n')

        summary.add_run('Kupiec检验结论: ').bold = True
        summary.add_run(f'{kupiec["conclusion"]}\n')

        # 添加分页符（除了最后一个）
        if alpha != list(results_dict.keys())[-1]:
            doc.add_page_break()

    # 添加总结部分
    doc.add_heading('总体总结', level=1)

    # 创建比较表格
    comparison_table = doc.add_table(rows=len(results_dict) + 1, cols=7)
    comparison_table.style = 'Light Shading'

    # 表头
    header_cells = comparison_table.rows[0].cells
    headers = ['α', '测试集损失', 'DQ统计量', 'DQ p值', 'DQ通过', 'Kupiec失败率', 'Kupiec结论']
    for i, header in enumerate(headers):
        header_cells[i].text = header

    # 填充数据
    for idx, (alpha, results) in enumerate(results_dict.items(), 1):
        row_cells = comparison_table.rows[idx].cells
        row_cells[0].text = str(alpha)
        row_cells[1].text = f'{results["test_QL"]:.6f}'
        row_cells[2].text = f'{results["dq_statistic"]:.4f}'
        row_cells[3].text = f'{results["dq_pvalue"]:.4f}'
        # 计算DQ是否通过
        dq_pass = "是" if results['dq_pvalue'] > 0.05 else "否"
        row_cells[4].text = dq_pass
        row_cells[5].text = f'{results["kupiec_result"]["failure_rate"]:.4f}'
        row_cells[6].text = results['kupiec_result']['conclusion']

    # 保存文档
    doc_path = os.path.join(output_dir, 'results.docx')
    doc.save(doc_path)

    print(f"\n{'=' * 60}")
    print(f"所有结果已保存到: {doc_path}")
    print(f"{'=' * 60}")

    return doc_path


# 主执行部分
def main():
    """
    主函数：对所有alpha值进行训练和测试
    """
    print("开始多alpha值LSTM训练和测试...")

    # 存储所有结果
    all_results = {}

    # 对每个alpha进行训练和测试
    for alpha in alphas:
        results = train_and_evaluate_for_alpha(
            alpha=alpha,
            X_train=X_train,
            y_train=y_train,
            X_validation=X_validation,
            y_validation=y_validation,
            X_test=X_test,
            y_test=y_test,
            input_size=6,
            window_size=21
        )

        all_results[alpha] = results

        # 打印当前alpha的结果摘要
        print(f"\nAlpha={alpha} 结果摘要:")
        print(f"  测试集损失函数值: {results['test_QL']:.6f}")
        print(f"  DQ检验: 统计量={results['dq_statistic']:.4f}, p值={results['dq_pvalue']:.4f}")
        print(f"  Kupiec检验: 失败率={results['kupiec_result']['failure_rate']:.4f}, "
              f"结论={results['kupiec_result']['conclusion']}")

    # 创建可视化并保存到Word文档
    doc_path = create_visualizations_and_save(all_results)

    # 在控制台显示最终比较
    print("\n" + "=" * 70)
    print("不同α值的模型性能比较:")
    print("=" * 70)
    print(f"{'α':<8} {'测试集损失':<15} {'DQ p值':<10} {'DQ通过':<10} {'Kupiec失败率':<15} {'Kupiec结论':<12}")
    print("-" * 70)

    for alpha, results in all_results.items():
        dq_pass = "是" if results['dq_pvalue'] > 0.05 else "否"
        print(f"{alpha:<8} {results['test_QL']:<15.6f} {results['dq_pvalue']:<10.4f} "
              f"{dq_pass:<10} {results['kupiec_result']['failure_rate']:<15.4f} "
              f"{results['kupiec_result']['conclusion']:<12}")

    return all_results, doc_path


# 执行主函数
if __name__ == "__main__":
    all_results, doc_path = main()